import os
import logging
import tempfile
import traceback
from datetime import datetime
from typing import Optional, Dict, Any, Tuple
from flask import Flask, request, jsonify, send_from_directory
from flask_cors import CORS
from werkzeug.utils import secure_filename
from werkzeug.exceptions import RequestEntityTooLarge
import requests
from pathlib import Path
from PIL import Image
import io
from dotenv import load_dotenv
load_dotenv() 
try:
    from PyPDF2 import PdfReader
    PDF_SUPPORT = True
    print("✅ PDF support enabled")
except ImportError:
    print("⚠️ Warning: PyPDF2 not installed. PDF support disabled.")
    PDF_SUPPORT = False

try:
    from docx import Document
    DOCX_SUPPORT = True
    print("✅ DOCX support enabled")
except ImportError:
    print("⚠️ Warning: python-docx not installed. DOCX support disabled.")
    DOCX_SUPPORT = False

try:
    import pytesseract
    from PIL import Image, ImageEnhance, ImageFilter
    OCR_SUPPORT = True
    print("✅ OCR support enabled with Tesseract")
except ImportError:
    print("⚠️ Warning: pytesseract or PIL not installed. Image OCR support disabled.")
    OCR_SUPPORT = False

logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)

app = Flask(__name__)
CORS(app)

class Config:
    GROQ_API_KEY = os.getenv("GROQ_API_KEY")
    GROQ_API_URL = "https://api.groq.com/openai/v1/chat/completions"
    MAX_CONTENT_LENGTH = 16 * 1024 * 1024
    UPLOAD_FOLDER = tempfile.gettempdir()
    
    ALLOWED_EXTENSIONS = {'txt'}
    if PDF_SUPPORT:
        ALLOWED_EXTENSIONS.add('pdf')
    if DOCX_SUPPORT:
        ALLOWED_EXTENSIONS.update({'docx', 'doc'})
    if OCR_SUPPORT:
        ALLOWED_EXTENSIONS.update({'png', 'jpg', 'jpeg', 'gif', 'bmp', 'tiff', 'webp'})

app.config.from_object(Config)

class DocumentAnalyzer:
    
    @staticmethod
    def is_allowed_file(filename: str) -> bool:
        if not filename or '.' not in filename:
            return False
        extension = filename.rsplit('.', 1)[1].lower()
        return extension in Config.ALLOWED_EXTENSIONS
    
    @staticmethod
    def extract_text_from_file(file) -> Tuple[str, Optional[str]]:
        try:
            filename = secure_filename(file.filename).lower()
            logger.info(f"Processing file: {filename}")
            
            if filename.endswith('.pdf'):
                return DocumentAnalyzer._extract_pdf_text(file), None
            elif filename.endswith(('.docx', '.doc')):
                return DocumentAnalyzer._extract_docx_text(file), None
            elif filename.endswith('.txt'):
                return DocumentAnalyzer._extract_txt_text(file), None
            elif filename.endswith(('.png', '.jpg', '.jpeg', '.gif', '.bmp', '.tiff', '.webp')):
                return DocumentAnalyzer._extract_image_text(file), None
            else:
                return "", "Unsupported file format"
                
        except Exception as e:
            logger.error(f"Error extracting text from {filename}: {str(e)}")
            return "", f"Error processing file: {str(e)}"
    
    @staticmethod
    def _extract_pdf_text(file) -> str:
        if not PDF_SUPPORT:
            raise Exception("PDF support not available")
        
        try:
            reader = PdfReader(file)
            text_content = []
            
            for page_num, page in enumerate(reader.pages):
                try:
                    page_text = page.extract_text() or ""
                    if page_text.strip():
                        text_content.append(page_text)
                except Exception as e:
                    logger.warning(f"Error reading page {page_num + 1}: {str(e)}")
                    continue
            
            extracted_text = "\n".join(text_content)
            return extracted_text
            
        except Exception as e:
            raise Exception(f"PDF processing error: {str(e)}")
    
    @staticmethod
    def _extract_docx_text(file) -> str:
        if not DOCX_SUPPORT:
            raise Exception("DOCX support not available")
        
        try:
            doc = Document(file)
            paragraphs = []
            
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    paragraphs.append(text)
            
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text = cell.text.strip()
                        if text:
                            paragraphs.append(text)
            
            extracted_text = "\n".join(paragraphs)
            return extracted_text
            
        except Exception as e:
            raise Exception(f"DOCX processing error: {str(e)}")
    
    @staticmethod
    def _extract_txt_text(file) -> str:
        try:
            content = file.read()
            
            for encoding in ['utf-8', 'utf-16', 'latin-1', 'cp1252']:
                try:
                    text = content.decode(encoding)
                    return text
                except UnicodeDecodeError:
                    continue
            
            text = content.decode('utf-8', errors='ignore')
            return text
            
        except Exception as e:
            raise Exception(f"TXT processing error: {str(e)}")
    
    @staticmethod
    def _extract_image_text(file) -> str:
        if not OCR_SUPPORT:
            raise Exception("OCR support not available")
        
        try:
            image = Image.open(file)
            
            if image.mode != 'RGB':
                image = image.convert('RGB')
            
            image = DocumentAnalyzer._preprocess_image_for_ocr(image)
            extracted_text = DocumentAnalyzer._ocr_with_fallback(image)
            cleaned_text = DocumentAnalyzer._clean_ocr_text(extracted_text)
            
            return cleaned_text
            
        except Exception as e:
            raise Exception(f"Image OCR processing error: {str(e)}")
    
    @staticmethod
    def _preprocess_image_for_ocr(image: Image.Image) -> Image.Image:
        try:
            width, height = image.size
            if width < 1000 or height < 1000:
                scale_factor = max(1200/width, 1200/height)
                new_width = int(width * scale_factor)
                new_height = int(height * scale_factor)
                image = image.resize((new_width, new_height), Image.Resampling.LANCZOS)
            
            enhancer = ImageEnhance.Contrast(image)
            image = enhancer.enhance(1.2)
            
            enhancer = ImageEnhance.Sharpness(image)
            image = enhancer.enhance(1.1)
            
            image = image.filter(ImageFilter.MedianFilter(size=1))
            
            return image
            
        except Exception as e:
            logger.warning(f"Image preprocessing failed: {e}")
            return image
    
    @staticmethod
    def _ocr_with_fallback(image: Image.Image) -> str:
        configs = [
            '--psm 6 --oem 3',
            '--psm 4 --oem 3',
            '--psm 3 --oem 3',
            '--psm 1 --oem 3',
            '--psm 6 --oem 1',
        ]
        
        best_result = ""
        max_length = 0
        
        for config in configs:
            try:
                result = pytesseract.image_to_string(image, config=config)
                if len(result) > max_length:
                    max_length = len(result)
                    best_result = result
            except Exception as e:
                continue
        
        if not best_result:
            try:
                best_result = pytesseract.image_to_string(image)
            except Exception as e:
                raise Exception("OCR extraction failed")
        
        return best_result
    
    @staticmethod
    def _clean_ocr_text(text: str) -> str:
        if not text:
            return ""
        
        lines = text.split('\n')
        cleaned_lines = []
        
        for line in lines:
            line = line.strip()
            
            if len(line) < 2:
                continue
                
            alphanumeric_chars = sum(1 for c in line if c.isalnum())
            if alphanumeric_chars < len(line) * 0.3 and len(line) > 5:
                continue
            
            cleaned_lines.append(line)
        
        cleaned_text = '\n'.join(cleaned_lines)
        
        import re
        cleaned_text = re.sub(r'\n\s*\n', '\n\n', cleaned_text)
        cleaned_text = re.sub(r' +', ' ', cleaned_text)
        
        return cleaned_text.strip()

class GroqAnalyzer:
    
    @staticmethod
    def get_analysis_prompt(text: str, analysis_type: str) -> str:
        
        prompts = {
            "resume": f"""You are an experienced career counselor and resume expert with expertise in computer science and technology careers.

Analyze the following resume comprehensively and provide detailed, actionable feedback. Structure your response with clear sections:

**STRENGTHS:**
- Key qualifications and standout features
- Well-presented skills and experiences  
- Notable achievements or projects
- Technical competencies demonstrated

**AREAS FOR IMPROVEMENT:**
- Missing critical elements (skills, experiences, formatting)
- Weak sections that need strengthening
- Suggestions for better presentation
- Industry-specific recommendations

**SKILLS ANALYSIS:**
- Technical skills identified and their relevance
- Programming languages and technologies
- Soft skills demonstrated through experiences
- Skills gaps for target roles in tech industry

**PROJECT & EXPERIENCE EVALUATION:**
- Quality and relevance of projects described
- Professional experience assessment
- Academic projects and their industry applicability
- Recommendations for portfolio enhancement

**CAREER GUIDANCE:**
- Suitable career paths based on background
- Industry recommendations (web dev, data science, AI/ML, etc.)
- Next steps for career development
- Certification or learning recommendations

**OVERALL ASSESSMENT:**
- Resume strength rating (1-10) with justification
- Key recommendations for immediate improvement
- Long-term career development advice

Resume Content:
{text}""",

            "marks_card": f"""You are an academic advisor and career counselor specializing in computer science education and student performance analysis.

Analyze the following academic transcript/marks card comprehensively and provide detailed insights:

**ACADEMIC PERFORMANCE:**
- Overall GPA/percentage analysis with context
- Grade distribution and performance patterns
- Semester-wise performance trends and consistency
- Comparison with typical CS program expectations

**SUBJECT ANALYSIS:**
- Strongest performing subjects and their significance
- Subjects needing improvement and strategies
- Core CS subjects vs electives performance comparison
- Technical vs theoretical subject performance

**SKILL ASSESSMENT:**
- Technical competencies indicated by coursework performance
- Programming and software development capabilities
- Mathematical and analytical problem-solving skills
- Areas of academic strength and specialization potential

**CAREER READINESS:**
- Industry readiness based on academic performance
- Alignment with different CS career paths
- Graduate school readiness assessment
- Professional skill development recommendations

**RECOMMENDATIONS:**
- Subjects/areas requiring focused attention and study strategies
- Career paths best aligned with academic strengths
- Further education, certification, or skill development suggestions
- Industry-specific preparation recommendations

**IMPROVEMENT STRATEGY:**
- Specific study recommendations for weak areas
- Skill development priorities for career goals
- Academic goal setting and achievement strategies
- Resource recommendations for continued learning

**OVERALL ASSESSMENT:**
- Academic performance rating with detailed justification
- Key strengths and competitive advantages
- Critical areas for improvement and growth
- Personalized career development roadmap

Marks Card/Transcript Content:
{text}"""
        }
        
        return prompts.get(analysis_type, f"Analyze the following document comprehensively:\n{text}")
    
    @staticmethod
    def analyze_with_groq(text: str, analysis_type: str) -> Tuple[str, Optional[str]]:
        if not text.strip():
            return "", "No text content found in the document"
        
        if len(text) > 15000:
            text = text[:10000] + "\n[Content truncated for analysis...]\n" + text[-5000:]
        
        try:
            prompt = GroqAnalyzer.get_analysis_prompt(text, analysis_type)
            
            headers = {
                "Authorization": f"Bearer {Config.GROQ_API_KEY}",
                "Content-Type": "application/json"
            }
            
            data = {
                "model": "llama3-70b-8192",
                "messages": [
                    {
                        "role": "system", 
                        "content": "You are a helpful academic and career advisor with expertise in computer science and technology careers. Provide detailed, actionable insights in a well-structured format using markdown-style formatting for better readability."
                    },
                    {
                        "role": "user", 
                        "content": prompt
                    }
                ],
                "temperature": 0.7,
                "max_tokens": 3000,
                "top_p": 0.9
            }
            
            response = requests.post(
                Config.GROQ_API_URL, 
                headers=headers, 
                json=data,
                timeout=45
            )
            
            if response.status_code == 200:
                result = response.json()
                analysis_content = result['choices'][0]['message']['content']
                return analysis_content, None
            else:
                error_msg = f"Groq API Error {response.status_code}: {response.text}"
                return "", error_msg
                
        except requests.exceptions.Timeout:
            return "", "Analysis request timeout - please try again"
        except requests.exceptions.RequestException as e:
            return "", f"Network error during analysis: {str(e)}"
        except Exception as e:
            return "", f"Analysis error: {str(e)}"

@app.errorhandler(413)
@app.errorhandler(RequestEntityTooLarge)
def handle_file_too_large(e):
    return jsonify({
        "success": False,
        "error": "File too large",
        "message": "Maximum file size is 16MB"
    }), 413

@app.errorhandler(500)
def handle_internal_error(e):
    return jsonify({
        "success": False,
        "error": "Internal server error",
        "message": "Something went wrong processing your request"
    }), 500

@app.errorhandler(400)
def handle_bad_request(e):
    return jsonify({
        "success": False,
        "error": "Bad request",
        "message": str(e)
    }), 400

def validate_file_request():
    if 'file' not in request.files:
        return {"error": "No file provided", "success": False}, 400
    
    file = request.files['file']
    if file.filename == '':
        return {"error": "No file selected", "success": False}, 400
    
    if not DocumentAnalyzer.is_allowed_file(file.filename):
        return {
            "error": "Invalid file type",
            "message": f"Allowed formats: {', '.join(Config.ALLOWED_EXTENSIONS)}",
            "success": False
        }, 400
    
    return None, None

@app.route('/', methods=['GET'])
def index():
    return jsonify({
        "service": "Student Document Analysis API",
        "version": "2.0.0",
        "status": "online",
        "supported_formats": list(Config.ALLOWED_EXTENSIONS),
        "max_file_size": "16MB",
        "features": {
            "pdf_support": PDF_SUPPORT,
            "docx_support": DOCX_SUPPORT,
            "ocr_support": OCR_SUPPORT
        },
        "endpoints": {
            "health": "GET /health",
            "resume_analysis": "POST /upload-resume",
            "marks_analysis": "POST /upload-marks-card",
            "generic_analysis": "POST /analyze"
        }
    })

@app.route('/health', methods=['GET'])
def health_check():
    return jsonify({
        "status": "healthy",
        "timestamp": str(datetime.now()),
        "api_available": True,
        "groq_configured": bool(Config.GROQ_API_KEY and Config.GROQ_API_KEY != "your-groq-api-key"),
        "features": {
            "pdf_support": PDF_SUPPORT,
            "docx_support": DOCX_SUPPORT,
            "ocr_support": OCR_SUPPORT
        },
        "supported_formats": list(Config.ALLOWED_EXTENSIONS)
    })

@app.route('/upload-resume', methods=['POST'])
def upload_resume():
    error_response, status_code = validate_file_request()
    if error_response:
        return jsonify(error_response), status_code
    
    file = request.files['file']
    
    try:
        text_content, extract_error = DocumentAnalyzer.extract_text_from_file(file)
        
        if extract_error:
            return jsonify({
                "success": False,
                "error": extract_error
            }), 400
        
        if not text_content.strip():
            return jsonify({
                "success": False,
                "error": "No readable text found in the document"
            }), 400
        
        analysis, analysis_error = GroqAnalyzer.analyze_with_groq(text_content, "resume")
        
        if analysis_error:
            return jsonify({
                "success": False,
                "error": analysis_error
            }), 500
        
        return jsonify({
            "success": True,
            "filename": secure_filename(file.filename),
            "analysis": analysis,
            "document_type": "resume",
            "extracted_text_length": len(text_content)
        })
        
    except Exception as e:
        return jsonify({
            "success": False,
            "error": "Processing failed",
            "message": str(e)
        }), 500

@app.route('/upload-marks-card', methods=['POST'])
def upload_marks_card():
    error_response, status_code = validate_file_request()
    if error_response:
        return jsonify(error_response), status_code
    
    file = request.files['file']
    
    try:
        text_content, extract_error = DocumentAnalyzer.extract_text_from_file(file)
        
        if extract_error:
            return jsonify({
                "success": False,
                "error": extract_error
            }), 400
        
        if not text_content.strip():
            return jsonify({
                "success": False,
                "error": "No readable text found in the document"
            }), 400
        
        analysis, analysis_error = GroqAnalyzer.analyze_with_groq(text_content, "marks_card")
        
        if analysis_error:
            return jsonify({
                "success": False,
                "error": analysis_error
            }), 500
        
        return jsonify({
            "success": True,
            "filename": secure_filename(file.filename),
            "analysis": analysis,
            "document_type": "marks_card",
            "extracted_text_length": len(text_content)
        })
        
    except Exception as e:
        return jsonify({
            "success": False,
            "error": "Processing failed",
            "message": str(e)
        }), 500

@app.route('/analyze', methods=['POST'])
def generic_analyze():
    error_response, status_code = validate_file_request()
    if error_response:
        return jsonify(error_response), status_code
    
    file = request.files['file']
    analysis_type = request.form.get('type', 'resume').lower()
    
    if analysis_type not in ['resume', 'marks_card']:
        return jsonify({
            "success": False,
            "error": "Invalid analysis type",
            "message": "Type must be 'resume' or 'marks_card'"
        }), 400
    
    try:
        text_content, extract_error = DocumentAnalyzer.extract_text_from_file(file)
        if extract_error:
            return jsonify({
                "success": False,
                "error": extract_error
            }), 400
        
        if not text_content.strip():
            return jsonify({
                "success": False,
                "error": "No readable text found in the document"
            }), 400
        
        analysis, analysis_error = GroqAnalyzer.analyze_with_groq(text_content, analysis_type)
        if analysis_error:
            return jsonify({
                "success": False,
                "error": analysis_error
            }), 500
        
        return jsonify({
            "success": True,
            "filename": secure_filename(file.filename),
            "analysis": analysis,
            "document_type": analysis_type,
            "extracted_text_length": len(text_content)
        })
        
    except Exception as e:
        return jsonify({
            "success": False,
            "error": "Processing failed",
            "message": str(e)
        }), 500

@app.route('/supported-formats', methods=['GET'])
def get_supported_formats():
    return jsonify({
        "supported_formats": list(Config.ALLOWED_EXTENSIONS),
        "features": {
            "pdf_support": PDF_SUPPORT,
            "docx_support": DOCX_SUPPORT,
            "ocr_support": OCR_SUPPORT
        }
    })

if __name__ == '__main__':
    if not Config.GROQ_API_KEY:
        print("⚠️ WARNING: GROQ_API_KEY not found in environment variables")

    # Run Flask app with port 10000 (required by Render Docker)
    app.run(
        debug=os.getenv("FLASK_DEBUG", "False").lower() == "true",
        host='0.0.0.0',
        port=int(os.getenv("PORT", 10000))
    )